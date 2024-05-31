from docx import Document

# Create a new Document
doc = Document()

# Q1) Operations on a NumPy Array
doc.add_heading('Q1) Operations on a NumPy Array', level=1)
q1_code = """
import numpy as np

# Initialize a 3x3 NumPy array with integer values
array = np.array([[1, 2, 3], [4, 5, 6], [7, 8, 9]])

# Multiply the entire array by 2
multiplied_array = array * 2

# Add 5 to each element of the array
added_array = multiplied_array + 5

# Calculate the square of each element in the array
squared_array = np.square(added_array)

# Print the original array and the results of each operation
print("Original Array:")
print(array)
print("\\nArray after multiplying by 2:")
print(multiplied_array)
print("\\nArray after adding 5:")
print(added_array)
print("\\nArray after squaring each element:")
print(squared_array)
"""
doc.add_paragraph(q1_code)

# Q2) Slicing Operations on a NumPy Array
doc.add_heading('Q2) Slicing Operations on a NumPy Array', level=1)
q2_code = """
import numpy as np

# Initialize a 3x3 NumPy array with integer values
array = np.array([[1, 2, 3], [4, 5, 6], [7, 8, 9]])

# Extract the first row of the array
first_row = array[0, :]

# Extract the last column of the array
last_column = array[:, -1]

# Extract a 2x2 sub-array from the center of the original array
center_subarray = array[1:3, 1:3]

# Print the results of the slicing operations
print("Original Array:")
print(array)
print("\\nFirst row of the array:")
print(first_row)
print("\\nLast column of the array:")
print(last_column)
print("\\n2x2 sub-array from the center:")
print(center_subarray)
"""
doc.add_paragraph(q2_code)

# Q3) DataFrame for Students' Names and Marks
doc.add_heading('Q3) DataFrame for Students\' Names and Marks', level=1)
q3_code = """
import pandas as pd

# Create a DataFrame to store names and marks of 10 students
data = {
    'Name': ['Alice', 'Bob', 'Charlie', 'David', 'Eva', 'Frank', 'Grace', 'Hannah', 'Ivan', 'Jack'],
    'Marks': [85, 92, 78, 90, 88, 76, 95, 89, 77, 84]
}

df_students = pd.DataFrame(data)

# Print the DataFrame
print("DataFrame of Students' Names and Marks:")
print(df_students)
"""
doc.add_paragraph(q3_code)

# Q4) DataFrame for Employees' Names and Incomes
doc.add_heading('Q4) DataFrame for Employees\' Names and Incomes', level=1)
q4_code = """
import pandas as pd

# Create a DataFrame to store names and incomes of 5 employees
data = {
    'Employee_name': ['John', 'Emma', 'Robert', 'Sophia', 'Michael'],
    'Income': [70000, 80000, 75000, 82000, 78000]
}

df_employees = pd.DataFrame(data, index=['a', 'b', 'c', 'd', 'e'])

# Print the DataFrame
print("DataFrame of Employees' Names and Incomes:")
print(df_employees)
"""
doc.add_paragraph(q4_code)

# Q5) Bar Plot for Frequency of Occurrences
doc.add_heading('Q5) Bar Plot for Frequency of Occurrences', level=1)
q5_code = """
import matplotlib.pyplot as plt

# Dataset representing the frequency of occurrences
x = ['A', 'B', 'C', 'D', 'E']
y = [10, 20, 15, 25, 30]

# Create a bar plot
plt.bar(x, y, color='blue')

# Set titles and labels
plt.title('Frequency of Occurrences')
plt.xlabel('Categories')
plt.ylabel('Frequency')

# Display the plot
plt.show()
"""
doc.add_paragraph(q5_code)

# Save the document
doc.save('answers.docx')
